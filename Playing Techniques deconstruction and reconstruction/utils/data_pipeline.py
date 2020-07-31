import docx
import os 
import pandas as pd
from config import data_model, style_transfer
from music21 import *

key = ['A','B','C','D','E','F','G']
time = ['2/4','4/4','1/4']
X = [] 
Y = [] 
X_base = [] 
X_dic = [] 
k_dic = []
meter_dic = []
dur = ['z','x','c','a']
low = ['*','v','b','n']
high = ['8','*'] 
Note = ['0','1','q','2','w','3','4','r','5','T','6','y','7']
T = ['T','t','H','h','B','d','D','X','Y','y','Q','K','L','l','S','F','q','f','x']

key_list = ['C','G','D','A','E','B','F#','C#','F','Bb','Eb','Ab','Db','Gb','Cb']
p = ['C','C#','D','D#','E','F','F#','G','G#','A','A#','B']
pitch_list = []

for i in range(6):
    for j in range(12):
        st = p[j] + str(i+1)
        pitch_list.append(st)

def pro(t):
    state = []
    for i in range(len(t)):
        if (t[i] in T) == True:
            state.append(t[i])
    state.sort()
    state = ''.join(state)
    return state

def encode(x,y,z,flag):
    if (data_model.mode == 3) or (flag == 1):
        result = [str(x)+str(y),str(z)]  
    else:   
        result = str(x) + str(y) + str(z)
    return result

def encode2(x,y,flag):
    if (data_model.mode == 3) or (flag == 1):
        result = [str(x),str(y)]
    else:
        result = str(x) + str(x) + str(y)
    return result

def encode3(x,y,z):
    result = []
    for item in x:
        result.append(item)
    result.append(y)
    result.append(z)
    return result

key_list = ['C','G','D','A','E','B','F#','C#','F','Bb','Eb','Ab','Db','Gb','Cb']
p = ['C','C#','D','D#','E','F','F#','G','G#','A','A#','B']

def transform_pitch(ch,k,high_low):
    temp = k + str(high_low)
    pitc = 0 
    do_index = pitch_list.index(temp)
    if ch == '1':
        pitc = temp
    elif ch == 'q':
        pitc = pitch_list[do_index+1]        
    elif ch == '2':
        pitc = pitch_list[do_index+2]
    elif ch == 'w':
        pitc = pitch_list[do_index+3]
    elif ch == '3':
        pitc = pitch_list[do_index+4]
    elif ch =='4':
        pitc = pitch_list[do_index+5]
    elif ch == 'r':
        pitc = pitch_list[do_index+6]
    elif ch == '5':
        pitc = pitch_list[do_index+7]
    elif ch == 'T':
        pitc = pitch_list[do_index+8]
    elif ch == '6':
        pitc = pitch_list[do_index+9]
    elif ch == 'y':
        pitc = pitch_list[do_index+10]
    elif ch == '7':
        pitc = pitch_list[do_index+11]
    return pitc

def transform_note(l,k,flag):
    global n,h,n2
    if flag == 0:
        if l[0] == '8':
            if len(l) == 2:
                n = note.Note(transform_pitch(l[1],k,5), quarterLength=1)
            elif len(l) == 3:
                if l[2] == ':':   
                    n = note.Note(transform_pitch(l[1],k,5), quarterLength=2)
                elif l[2] == "'":
                    n = note.Note(transform_pitch(l[1],k,5), quarterLength=1.5)
            elif len(l) == 4:
                n = note.Note(transform_pitch(l[1],k,5), quarterLength=3)
            elif len(l) == 5:
                n = note.Note(transform_pitch(l[1],k,5), quarterLength=4)      

                
        if (l[0] in Note) == True :
            if len(l) == 1:
                if l[0] == '0':
                    n = note.Rest(quarterLength=1)
                else:        
                    n = note.Note(transform_pitch(l[0],k,4), quarterLength=1)
            elif len(l) == 2:
                if l[1] == ':':  
                    if l[0] == '0':
                        n = note.Rest(quarterLength=2)
                    else:   
                        n = note.Note(transform_pitch(l[0],k,4), quarterLength=2)
                elif l[1] == "'":
                    if l[0] == '0':
                        n = note.Rest(quarterLength=1.5)
                    else: 
                        n = note.Note(transform_pitch(l[0],k,4), quarterLength=1.5)
            elif len(l) == 3:
                if l[0] == '0':
                    n = note.Rest(quarterLength=3)
                else: 
                    n = note.Note(transform_pitch(l[0],k,4), quarterLength=3)
            elif len(l) == 4:
                if l[0] == '0':
                    n = note.Rest(quarterLength=4)
                else: 
                    n = note.Note(transform_pitch(l[0],k,4), quarterLength=4)
                
                
        elif (l[0] in dur) == True:
            if l[1] == '8':
                if l[0] == 'z':
                    if len(l) == 3:
                        n = note.Note(transform_pitch(l[2],k,5), quarterLength=0.5)
                    else:
                        n = note.Note(transform_pitch(l[2],k,5), quarterLength=0.75)                 
                elif l[0] == 'x':
                    if len(l) == 3:
                        n = note.Note(transform_pitch(l[2],k,5), quarterLength=0.25)
                    else:
                        if l[2] == 't' or l[2] == 's':
                            n = note.Note(transform_pitch(l[3],k,5), quarterLength=1/6)
                        elif l[2] == 'f':
                            n = note.Note(transform_pitch(l[3],k,5), quarterLength=1/5)
                        else:
                            n = note.Note(transform_pitch(l[2],k,5), quarterLength=0.375)               
                elif l[0] == 'c':
                    if len(l) == 3:
                        n = note.Note(transform_pitch(l[2],k,5), quarterLength=0.125)
                    else:
                        n = note.Note(transform_pitch(l[2],k,5), quarterLength=0.1875)
                elif l[0] == 'a':
                    if len(l) == 3:
                        n = note.Note(transform_pitch(l[2],k,5), quarterLength=0.0625)
                    else:
                        n = note.Note(transform_pitch(l[2],k,5), quarterLength=0.09375)
                        
            elif l[1] == '9':
                if l[0] == 'z':
                    if len(l) == 3:
                        n = note.Note(transform_pitch(l[2],k,6), quarterLength=0.5)
                    else:
                        n = note.Note(transform_pitch(l[2],k,6), quarterLength=0.75)                 
                elif l[0] == 'x':
                    if len(l) == 3:
                        n = note.Note(transform_pitch(l[2],k,6), quarterLength=0.25)
                    else:
                        if l[2] == 't' or l[2] == 's':
                            n = note.Note(transform_pitch(l[3],k,6), quarterLength=1/6)
                        elif l[2] == 'f':
                            n = note.Note(transform_pitch(l[3],k,6), quarterLength=1/5)
                        else:
                            n = note.Note(transform_pitch(l[2],k,6), quarterLength=0.375)               
                elif l[0] == 'c':
                    if len(l) == 3:
                        n = note.Note(transform_pitch(l[2],k,6), quarterLength=0.125)
                    else:
                        n = note.Note(transform_pitch(l[2],k,6), quarterLength=0.1875)
                elif l[0] == 'a':
                    if len(l) == 3:
                        n = note.Note(transform_pitch(l[2],k,6), quarterLength=0.0625)
                    else:
                        n = note.Note(transform_pitch(l[2],k,6), quarterLength=0.09375)
                                
            else:
                if l[0] == 'z':
                    if len(l) == 2:
                        if l[1] == 't' or l[1] == 's':
                            n = note.Note(transform_pitch(l[2],k,4), quarterLength=1/3)
                        if l[1] == '0':
                            n = note.Rest(quarterLength=0.5)
                        else:
                            n = note.Note(transform_pitch(l[1],k,4), quarterLength=0.5)
                    else:
                        if l[1] == '0':
                            n = note.Rest(quarterLength=0.75)
                        else:
                            n = note.Note(transform_pitch(l[1],k,4), quarterLength=0.75)                  
                elif l[0] == 'x':
                    if len(l) == 2:
                        if l[1] == '0':
                            n = note.Rest(quarterLength=0.25)
                        else:
                            n = note.Note(transform_pitch(l[1],k,4), quarterLength=0.25)
                    else:
                        if l[1] == 't' or l[1] == 's':
                            n = note.Note(transform_pitch(l[2],k,4), quarterLength=1/6)
                        elif l[1] == 'f':
                            n = note.Note(transform_pitch(l[2],k,4), quarterLength=1/5)
                        elif l[1] == '0':
                            n = note.Rest(quarterLength=0.375)
                        else:
                            n = note.Note(transform_pitch(l[1],k,4), quarterLength=0.375)                
                elif l[0] == 'c':
                    if len(l) == 2:
                        if l[1] == '0':
                            n = note.Rest(quarterLength=0.125)
                        else:
                            n = note.Note(transform_pitch(l[1],k,4), quarterLength=0.125)
                    else:
                        if l[1] == 't' or l[1] == 's':
                            n = note.Note(transform_pitch(l[2],k,4), quarterLength=1/12)
                        elif l[0] == '0':
                            n = note.Rest(quarterLength=0.1875)
                        else:
                            n = note.Note(transform_pitch(l[1],k,4), quarterLength=0.1875) 
                elif l[0] == 'a':
                    if len(l) == 2:
                        if l[1] == '0':
                            n = note.Rest(quarterLength=0.0625)
                        else:
                            n = note.Note(transform_pitch(l[1],k,4), quarterLength=0.0625)
                    else:
                        if l[0] == '0':
                            n = note.Rest(quarterLength=0.09375)
                        else:
                            n = note.Note(transform_pitch(l[1],k,4), quarterLength=0.09375)


        elif (l[0] in low) == True:
            if l[0] == '*':
                if len(l) == 2:
                    n = note.Note(transform_pitch(l[1],k,3), quarterLength=1)
                elif len(l) == 3:
                    if l[2] == ':':
                        n = note.Note(transform_pitch(l[1],k,3), quarterLength=2)
                    elif l[2] == "'":
                        n = note.Note(transform_pitch(l[1],k,3), quarterLength=1.5)
                elif len(l) == 4:
                    n = note.Note(transform_pitch(l[1],k,3), quarterLength=3)
                elif len(l) == 5:
                    n = note.Note(transform_pitch(l[1],k,3), quarterLength=4)
            elif l[0] == 'v':
                if len(l) == 2:
                    n = note.Note(transform_pitch(l[1],k,3), quarterLength=0.5) 
                else:
                    if l[1] == 't' or l[1] == 't':
                        n = note.Note(transform_pitch(l[2],k,3), quarterLength=1/3)
                    elif l[1] == 'f':
                        n = note.Note(transform_pitch(l[2],k,3), quarterLength=2/5)
                    else:
                        n = note.Note(transform_pitch(l[1],k,3), quarterLength=0.75)
            elif l[0] == 'b':
                if len(l) == 2:
                    n = note.Note(transform_pitch(l[1],k,3), quarterLength=0.25)  
                else:
                    if l[1] == 't' or l[1] == 't':
                        n = note.Note(transform_pitch(l[2],k,3), quarterLength=1/6)
                    elif l[1] == 'f':
                        n = note.Note(transform_pitch(l[2],k,3), quarterLength=1/5)
                    else:
                        n = note.Note(transform_pitch(l[1],k,3), quarterLength=0.375)
            elif l[0] == 'n':
                    n = note.Note(transform_pitch(l[1],k,3), quarterLength=0.125)
                    
        if n.name == 'rest':
            h = str(0) + str(n.quarterLength)
        else:
            h = str(n.name)+str(n.octave)+str(n.quarterLength)
        return h    

#------------------------------------------------------------------------------------------------------------------
    elif flag == 1:
        if l[0] == '8':
            if len(l) == 2:
                n2 = encode(l[1],5,1,flag)
            elif len(l) == 3:
                if l[2] == ':':   
                    n2 = encode(l[1],5,2,flag)
                elif l[2] == "'":
                    n2 = encode(l[1],5,1.5,flag)
            elif len(l) == 4:
                n2 = encode(l[1],5,3,flag)
            elif len(l) == 5:
                n2 = encode(l[1],5,4,flag)      

                
        if (l[0] in Note) == True :
            if len(l) == 1:
                if l[0] == '0':
                    n2 = encode2(0,1,flag)
                else:        
                    n2 = encode(l[0],4,1,flag)
            elif len(l) == 2:
                if l[1] == ':':  
                    if l[0] == '0':
                        n2 = encode2(0,2,flag)
                    else:   
                        n2 = encode(l[0],4,2,flag)
                elif l[1] == "'":
                    if l[0] == '0':
                        n2 = encode2(0,1.5,flag)
                    else: 
                        n2 = encode(l[0],4,1.5,flag)
            elif len(l) == 3:
                if l[0] == '0':
                    n2 = encode2(0,3,flag)
                else: 
                    n2 = encode(l[0],4,3,flag)
            elif len(l) == 4:
                if l[0] == '0':
                    n2 = encode2(0,4,flag)
                else: 
                    n2 = encode(l[0],4,4,flag)
                
                
        elif (l[0] in dur) == True:
            if l[1] == '8':
                if l[0] == 'z':
                    if len(l) == 3:
                        n2 = encode(l[2],5,0.5,flag)
                    else:
                        n2 = encode(l[2],5,0.75,flag)                
                elif l[0] == 'x':
                    if len(l) == 3:
                        n2 = encode(l[2],5,0.25,flag)
                    else:
                        if l[2] == 't' or l[2] == 's':
                            n2 = encode(l[3],5,1/6,flag)
                        elif l[2] == 'f':
                            n2 = encode(l[3],5,1/5,flag)
                        else:
                            n2 = encode(l[2],5,0.375,flag)              
                elif l[0] == 'c':
                    if len(l) == 3:
                        n2 = encode(l[2],5,0.125,flag)
                    else:
                        n2 = encode(l[2],5,0.1875,flag)
                elif l[0] == 'a':
                    if len(l) == 3:
                        n2 = encode(l[2],5,0.0625,flag)
                    else:
                        n2 = encode(l[2],5,0.09375,flag)
                        
            elif l[1] == '9':
                if l[0] == 'z':
                    if len(l) == 3:
                        n2 = encode(l[2],6,0.5,flag)
                    else:
                        n2 = encode(l[2],6,0.75,flag)                 
                elif l[0] == 'x':
                    if len(l) == 3:
                        n2 = encode(l[2],6,0.25,flag)
                    else:
                        if l[2] == 't' or l[2] == 's':
                            n2 = encode(l[3],6,1/6,flag)
                        elif l[2] == 'f':
                            n2 = encode(l[3],6,1/5,flag)
                        else:
                            n2 = encode(l[2],6,0.375,flag)               
                elif l[0] == 'c':
                    if len(l) == 3:
                        n2 = encode(l[2],6,0.125,flag)
                    else:
                        n2 = encode(l[2],6,0.1875,flag)
                elif l[0] == 'a':
                    if len(l) == 3:
                        n2 = encode(l[2],6,0.0625,flag)
                    else:
                        n2 = encode(l[2],6,0.09375,flag)
                                
            else:
                if l[0] == 'z':
                    if len(l) == 2:
                        if l[1] == 't' or l[1] == 's':
                            n2 = encode(l[2],4,1/3,flag)
                        if l[1] == '0':
                            n2 = encode2(0,0.5,flag)
                        else:
                            n2 = encode(l[1],4,0.5,flag)
                    else:
                        if l[1] == '0':
                            n2 = encode2(0,0.75,flag)
                        else:
                            n2 = encode(l[1],4,0.75,flag)                  
                elif l[0] == 'x':
                    if len(l) == 2:
                        if l[1] == '0':
                            n2 = encode2(0,0.25,flag)
                        else:
                            n2 = encode(l[1],4,0.25,flag)
                    else:
                        if l[1] == 't' or l[1] == 's':
                            n2 = encode(l[2],4,1/6,flag)
                        elif l[1] == 'f':
                            n2 = encode(l[2],4,1/5,flag)
                        elif l[1] == '0':
                            n2 = encode2(0,0.375,flag)
                        else:
                            n2 = encode(l[1],4,0.375,flag)                
                elif l[0] == 'c':
                    if len(l) == 2:
                        if l[1] == '0':
                            n2 = encode2(0,0.125,flag)
                        else:
                            n2 = encode(l[1],4,0.125,flag)
                    else:
                        if l[1] == 't' or l[1] == 's':
                            n2 = encode(l[2],4,1/12,flag)
                        elif l[0] == '0':
                            n2 = encode2(0,0.1875,flag)
                        else:
                            n2 = encode(l[1],4,0.1875,flag) 
                elif l[0] == 'a':
                    if len(l) == 2:
                        if l[1] == '0':
                            n2 = encode2(0,0.0625,flag)
                        else:
                            n2 = encode(l[1],4,0.0625,flag)
                    else:
                        if l[0] == '0':
                            n2 = encode2(0,0.09375,flag)
                        else:
                            n2 = encode(l[1],4,0.09375,flag)


        elif (l[0] in low) == True:
            if l[0] == '*':
                if len(l) == 2:
                    n2 = encode(l[1],3,1,flag)
                elif len(l) == 3:
                    if l[2] == ':':
                        n2 = encode(l[1],3,2,flag)
                    elif l[2] == "'":
                        n2 = encode(l[1],3,1.5,flag)
                elif len(l) == 4:
                    n2 = encode(l[1],3,3,flag)
                elif len(l) == 5:
                    n2 = encode(l[1],3,4,flag)
            elif l[0] == 'v':
                if len(l) == 2:
                    n2 = encode(l[1],3,0.5,flag) 
                else:
                    if l[1] == 't' or l[1] == 't':
                        n2 = encode(l[2],3,1/3,flag)
                    elif l[1] == 'f':
                        n2 = encode(l[2],3,2/5,flag)
                    else:
                        n2 = encode(l[1],3,0.75,flag)
            elif l[0] == 'b':
                if len(l) == 2:
                    n2 = encode(l[1],3,0.25,flag)  
                else:
                    if l[1] == 't' or l[1] == 't':
                        n2 = encode(l[2],3,1/6,flag)
                    elif l[1] == 'f':
                        n2 = encode(l[2],3,1/5,flag)
                    else:
                        n2 = encode(l[1],3,0.375,flag)
            elif l[0] == 'n':
                    n2 = encode(l[1],3,0.125,flag)
        return n2    
  

def cal_techique(l,k,meter,flag):
    tec = []
    no = []
    mu = 0
    tech = 0
    for ch in l:
        if (ch == '|') and (mu == 0):
            mu = 1
        elif (ch == '|') and (mu == 1):
            mu = 0
        elif mu == 1:
            tec.append(ch)
        elif mu == 0:
            no.append(ch)
    tec = ''.join(tec)
    no = ''.join(no)
    n = no
    no = transform_note(no,k,flag)
    for c in tec:
        if (c in T) == True:
            tech += 1

    if tech == 0:
        sta = '0'
    else:
        if data_model.mode !=3:
            sta = pro(tec)
        else:
            sta = tec
    return n,no,sta
    
        
def ifend(section, count):
    if count == len(section) - 1:
        return True
    elif (section[count] in Note) == True:
        if (count + 2) <= (len(section) - 1):
            if section[count+1] != ':' and section[count+1] != "'" and section[count+2] != ';':
                return True
            else:
                return False
        elif (count + 1) <= (len(section) - 1):
            if section[count+1] != ':' and section[count+1] != "'":
                return True
            else: 
                return False
        else:
            return False
    elif section[count] == "'" or section[count] == ";":
        return True
    elif (section[count] == ":") and (section[count+1] != ":"):
        return True
    else:
        return False


def text_process(text,k,meter,flag,X,Y):
    text = text.replace(" ","")
    section_list = text.split('/')       
    for section in section_list:
        count = 0
        mutex = 0
        current_note = []
        x_list = []
        x1_list = []
        y_list = []
        for character in section:
            if (character == '|') and (mutex == 0): 
                mutex = 1
                current_note.append(character)
            elif (character == '|') and (mutex == 1):
                mutex = 0
                current_note.append(character)
            elif (ifend(section,count) == True) and (mutex == 0):
                current_note.append(character)
                current_note = ''.join(current_note)
#                print(count,section)
#                print(current_note)
                x1,x,y = cal_techique(current_note,k,meter,flag)
                if flag == 0:
                    X_dic.append(x)
                y = ''.join(y)
                x_list.append(x)
                x1_list.append(x1)
                y_list.append(y)
                current_note = []
            else:
                current_note.append(character)
            count += 1
            if count == len(section):
                current_note = ''.join(current_note)
        X_base.append(x1_list)
        X.append(x_list)
        Y.append(y_list)

def pro_data(flag):
    X = [] 
    Y = [] 
    if data_model.mode == 1 or data_model.mode == 3: 
        path = data_model.data_path
        for root, dirs, files in os.walk(path):
            for file in files:
                f = docx.Document(os.path.join(path,file))
                chapter_number = int ((len(f.paragraphs) - 1) / 2)
                for i in range(chapter_number):
                    l = f.paragraphs[1+2*i].text.split(' ')
                    k = l[0][2:3]
                    meter = l[1]
                    text_process(f.paragraphs[2+2*i].text,k,meter,flag,X,Y)
    elif data_model.mode == 2:
        path = style_transfer.base_path
        f = docx.Document(path)
        chapter_number = int ((len(f.paragraphs) - 1) / 2)
        for i in range(chapter_number):
            l = f.paragraphs[1+2*i].text.split(' ')
            k = l[0][2:3]
            meter = l[1]
            text_process(f.paragraphs[2+2*i].text,k,meter,flag,X,Y)

    return X, Y


X,Y = pro_data(0) 
X2, Y2 = pro_data(1)


