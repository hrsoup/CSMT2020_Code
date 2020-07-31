import docx
from music21 import*
import os

from config import style_transfer

file = docx.Document(style_transfer.base_path)

score_name = file.paragraphs[0].text
chapter_number = int ((len(file.paragraphs) - 1) / 2)

score = stream.Score()
part = stream.Part()
part.partName = 'Di'
score.insert(0, metadata.Metadata())
score.metadata.title = score_name
#-------------------------------------------------------------------#

dur = ['z','x','c','a']
low = ['*','v','b','n']
high = ['8','*'] 
Note = ['0','1','2','3','4','5','6','7','q','w','r','T','y']
T = ['T','t','H','h','B','d','D','X','Y','y','Q','K','L','l','S','F','q','f','x']


key_list = ['C','G','D','A','E','B','F#','C#','F','Bb','Eb','Ab','Db','Gb','Cb']
p = ['C','C#','D','D#','E','F','F#','G','G#','A','A#','B']
pitch_list = []

for i in range(6):
    for j in range(12):
        st = p[j] + str(i+1)
        pitch_list.append(st)

def pro(t):
    state = []
    for i in range(len(t)):
        if (t[i] in T) == True:
            state.append(t[i])
    state.sort()
    state = ''.join(state)
    return state
        
def transform_pitch(ch,k,high_low):
    temp = k + str(high_low)
    pitc = 0 
    do_index = pitch_list.index(temp)
    if ch == '1':
        pitc = temp
    elif ch == 'q':
        pitc = pitch_list[do_index+1]        
    elif ch == '2':
        pitc = pitch_list[do_index+2]
    elif ch == 'w':
        pitc = pitch_list[do_index+3]
    elif ch == '3':
        pitc = pitch_list[do_index+4]
    elif ch =='4':
        pitc = pitch_list[do_index+5]
    elif ch == 'r':
        pitc = pitch_list[do_index+6]
    elif ch == '5':
        pitc = pitch_list[do_index+7]
    elif ch == 'T':
        pitc = pitch_list[do_index+8]
    elif ch == '6':
        pitc = pitch_list[do_index+9]
    elif ch == 'y':
        pitc = pitch_list[do_index+10]
    elif ch == '7':
        pitc = pitch_list[do_index+11]
    return pitc


def transform_note(l,t,stream1,k, pred1, pred2 ,row, co):
    print(l)
    if l[0] == '8':
        if len(l) == 2:
            n = note.Note(transform_pitch(l[1],k,5), quarterLength=1)
        elif len(l) == 3:
            if l[2] == ':':   
                n = note.Note(transform_pitch(l[1],k,5), quarterLength=2)
            elif l[2] == "'":
                n = note.Note(transform_pitch(l[1],k,5), quarterLength=1.5)
        elif len(l) == 4:
            n = note.Note(transform_pitch(l[1],k,5), quarterLength=3)
        elif len(l) == 5:
            n = note.Note(transform_pitch(l[1],k,5), quarterLength=4)      

            
    if (l[0] in Note) == True :
        if len(l) == 1:
            if l[0] == '0':
                n = note.Rest(quarterLength=1)
            else:        
                n = note.Note(transform_pitch(l[0],k,4), quarterLength=1)
        elif len(l) == 2:
            if l[1] == ':':  
                if l[0] == '0':
                    n = note.Rest(quarterLength=2)
                else:   
                    n = note.Note(transform_pitch(l[0],k,4), quarterLength=2)
            elif l[1] == "'":
                if l[0] == '0':
                    n = note.Rest(quarterLength=1.5)
                else: 
                    n = note.Note(transform_pitch(l[0],k,4), quarterLength=1.5)
        elif len(l) == 3:
            if l[0] == '0':
                n = note.Rest(quarterLength=3)
            else: 
                n = note.Note(transform_pitch(l[0],k,4), quarterLength=3)
        elif len(l) == 4:
            if l[0] == '0':
                n = note.Rest(quarterLength=4)
            else: 
                n = note.Note(transform_pitch(l[0],k,4), quarterLength=4)
            
            
    elif (l[0] in dur) == True:
        if l[1] == '8':
            if l[0] == 'z':
                if len(l) == 3:
                    n = note.Note(transform_pitch(l[2],k,5), quarterLength=0.5)
                else:
                    n = note.Note(transform_pitch(l[2],k,5), quarterLength=0.75)                 
            elif l[0] == 'x':
                if len(l) == 3:
                    n = note.Note(transform_pitch(l[2],k,5), quarterLength=0.25)
                else:
                    if l[2] == 't' or l[2] == 's':
                        n = note.Note(transform_pitch(l[3],k,5), quarterLength=1/6)
                    elif l[2] == 'f':
                        n = note.Note(transform_pitch(l[3],k,5), quarterLength=1/5)
                    else:
                        n = note.Note(transform_pitch(l[2],k,5), quarterLength=0.375)               
            elif l[0] == 'c':
                if len(l) == 3:
                    n = note.Note(transform_pitch(l[2],k,5), quarterLength=0.125)
                else:
                    n = note.Note(transform_pitch(l[2],k,5), quarterLength=0.1875)
            elif l[0] == 'a':
                if len(l) == 3:
                    n = note.Note(transform_pitch(l[2],k,5), quarterLength=0.0625)
                else:
                    n = note.Note(transform_pitch(l[2],k,5), quarterLength=0.09375)
                    
        elif l[1] == '9':
            if l[0] == 'z':
                if len(l) == 3:
                    n = note.Note(transform_pitch(l[2],k,6), quarterLength=0.5)
                else:
                    n = note.Note(transform_pitch(l[2],k,6), quarterLength=0.75)                 
            elif l[0] == 'x':
                if len(l) == 3:
                    n = note.Note(transform_pitch(l[2],k,6), quarterLength=0.25)
                else:
                    if l[2] == 't' or l[2] == 's':
                        n = note.Note(transform_pitch(l[3],k,6), quarterLength=1/6)
                    elif l[2] == 'f':
                        n = note.Note(transform_pitch(l[3],k,6), quarterLength=1/5)
                    else:
                        n = note.Note(transform_pitch(l[2],k,6), quarterLength=0.375)               
            elif l[0] == 'c':
                if len(l) == 3:
                    n = note.Note(transform_pitch(l[2],k,6), quarterLength=0.125)
                else:
                    n = note.Note(transform_pitch(l[2],k,6), quarterLength=0.1875)
            elif l[0] == 'a':
                if len(l) == 3:
                    n = note.Note(transform_pitch(l[2],k,6), quarterLength=0.0625)
                else:
                    n = note.Note(transform_pitch(l[2],k,6), quarterLength=0.09375)
                            
        else:
            if l[0] == 'z':
                if len(l) == 2:
                    if l[1] == 't' or l[1] == 's':
                        n = note.Note(transform_pitch(l[2],k,4), quarterLength=1/3)
                    if l[1] == '0':
                        n = note.Rest(quarterLength=0.5)
                    else:
                        n = note.Note(transform_pitch(l[1],k,4), quarterLength=0.5)
                else:
                    if l[1] == '0':
                        n = note.Rest(quarterLength=0.75)
                    else:
                        n = note.Note(transform_pitch(l[1],k,4), quarterLength=0.75)                  
            elif l[0] == 'x':
                if len(l) == 2:
                    if l[1] == '0':
                        n = note.Rest(quarterLength=0.25)
                    else:
                        n = note.Note(transform_pitch(l[1],k,4), quarterLength=0.25)
                else:
                    if l[1] == 't' or l[1] == 's':
                        n = note.Note(transform_pitch(l[2],k,4), quarterLength=1/6)
                    elif l[1] == 'f':
                        n = note.Note(transform_pitch(l[2],k,4), quarterLength=1/5)
                    elif l[1] == '0':
                        n = note.Rest(quarterLength=0.375)
                    else:
                        n = note.Note(transform_pitch(l[1],k,4), quarterLength=0.375)                
            elif l[0] == 'c':
                if len(l) == 2:
                    if l[1] == '0':
                        n = note.Rest(quarterLength=0.125)
                    else:
                        n = note.Note(transform_pitch(l[1],k,4), quarterLength=0.125)
                else:
                    if l[1] == 't' or l[1] == 's':
                        n = note.Note(transform_pitch(l[2],k,4), quarterLength=1/12)
                    elif l[0] == '0':
                        n = note.Rest(quarterLength=0.1875)
                    else:
                        n = note.Note(transform_pitch(l[1],k,4), quarterLength=0.1875) 
            elif l[0] == 'a':
                if len(l) == 2:
                    if l[1] == '0':
                        n = note.Rest(quarterLength=0.0625)
                    else:
                        n = note.Note(transform_pitch(l[1],k,4), quarterLength=0.0625)
                else:
                    if l[0] == '0':
                        n = note.Rest(quarterLength=0.09375)
                    else:
                        n = note.Note(transform_pitch(l[1],k,4), quarterLength=0.09375)


    elif (l[0] in low) == True:
        if l[0] == '*':
            if len(l) == 2:
                n = note.Note(transform_pitch(l[1],k,3), quarterLength=1)
            elif len(l) == 3:
                if l[2] == ':':
                    n = note.Note(transform_pitch(l[1],k,3), quarterLength=2)
                elif l[2] == "'":
                    n = note.Note(transform_pitch(l[1],k,3), quarterLength=1.5)
            elif len(l) == 4:
                n = note.Note(transform_pitch(l[1],k,3), quarterLength=3)
            elif len(l) == 5:
                n = note.Note(transform_pitch(l[1],k,3), quarterLength=4)
        elif l[0] == 'v':
            if len(l) == 2:
                n = note.Note(transform_pitch(l[1],k,3), quarterLength=0.5) 
            else:
                if l[1] == 't' or l[1] == 't':
                    n = note.Note(transform_pitch(l[2],k,3), quarterLength=1/3)
                elif l[1] == 'f':
                    n = note.Note(transform_pitch(l[2],k,3), quarterLength=2/5)
                else:
                    n = note.Note(transform_pitch(l[1],k,3), quarterLength=0.75)
        elif l[0] == 'b':
            if len(l) == 2:
                n = note.Note(transform_pitch(l[1],k,3), quarterLength=0.25)  
            else:
                if l[1] == 't' or l[1] == 't':
                    n = note.Note(transform_pitch(l[2],k,3), quarterLength=1/6)
                elif l[1] == 'f':
                    n = note.Note(transform_pitch(l[2],k,3), quarterLength=1/5)
                else:
                    n = note.Note(transform_pitch(l[1],k,3), quarterLength=0.375)
        elif l[0] == 'n':
                n = note.Note(transform_pitch(l[1],k,3), quarterLength=0.125)
    n.addLyric(t)
    n.addLyric(pred1[row][co])
    n.addLyric(pred2[row][co])  
    co += 1 
    stream1.append(n)
    return stream1, row, co


def ifend(section, count):
    if count == len(section) - 1:
        return True
    elif (section[count] in Note) == True:
        if (count + 2) <= (len(section) - 1):
            if section[count+1] != ':' and section[count+1] != "'" and section[count+2] != ';':
                return True
            else:
                return False
        elif (count + 1) <= (len(section) - 1):
            if section[count+1] != ':' and section[count+1] != "'":
                return True
            else: 
                return False
        else:
            return False
    elif section[count] == "'" or section[count] == ";":
        return True
    elif (section[count] == ":") and (section[count+1] != ":"):
        return True
    else:
        return False

def cal_techique(l,k,meter):
    tec = []
    no = []
    mu = 0
    tech = 0
    for ch in l:
        if (ch == '|') and (mu == 0):
            mu = 1
        elif (ch == '|') and (mu == 1):
            mu = 0
        elif mu == 1:
            tec.append(ch)
        elif mu == 0:
            no.append(ch)
    tec = ''.join(tec)
    no = ''.join(no)
    
    for c in tec:
        if (c in T) == True:
            tech += 1

    if tech == 0:
        sta = '0'
    else:
        sta = tec
    return no,sta

def text_process(text, stream1,k, pred1, pred2,row, co):
    text = text.replace(" ","")
    section_list = text.split('/')          
    for section in section_list:
        count = 0
        mutex = 0
        current_note = []
        for character in section:        
            if (character == '|') and (mutex == 0): 
                mutex = 1
                current_note.append(character)
            elif (character == '|') and (mutex == 1):
                mutex = 0
                current_note.append(character)
            elif (ifend(section,count) == True) and (mutex == 0):
                current_note.append(character)
                current_note = ''.join(current_note)
                x,y = cal_techique(current_note,k,meter)
                stream1,row,co = transform_note(x,y,stream1,k, pred1, pred2,row, co)
                current_note = []
            else:
                current_note.append(character)
            count += 1
            if count == len(section):
                current_note = ''.join(current_note)
        row += 1
        co = 0
                
    return stream1, row, co

def w2staff(pred1, pred2):
    row = 0
    co = 0
    for i in range(chapter_number):
            stream1 = stream.Stream()
            l = file.paragraphs[1+2*i].text.split(' ')
            k = l[0][2:3]
            index = int(key_list.index(k))
            num = int(l[2])
            stream1.append(key.KeySignature(index))
            stream1.append(meter.TimeSignature(l[1]))
            stream1.append(tempo.MetronomeMark(number=num))
            stream1, row, co = text_process(file.paragraphs[2+2*i].text,stream1,k, pred1, pred2, row, co)
            part.append(stream1)

    score.insert(0, part)
    file_path = os.path.join('transfer_results', str(score_name)+ '.xml')
    score.write('xml',fp=file_path) 